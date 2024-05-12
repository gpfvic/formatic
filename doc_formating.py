#coding: utf8

import os, subprocess, docx
from docx import Document # type: ignore
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


fpath = os.path.join(os.pardir, "demo.docx")


## read from docx file
origin_doc = Document(fpath)

text = []
for para in origin_doc.paragraphs:
    text.append(para.text)
text = list(filter(None, text))

## format and save new docx file
doc = Document() 

# 正文标题
# doc.styles['Heading 1'].font.name = '方正小标宋_GBK'  # Set a default font for Latin text
doc.styles['Heading 1']._element.rPr.rFonts.set(qn('w:eastAsia'), "方正小标宋_GBK")


doc.styles['Heading 1'].font.size = Pt(22)
doc.styles['Heading 1'].font.color.rgb = RGBColor(0,0,0)


run = doc.add_heading("", level=1).add_run(text[0])
# 设置西文字体
run.font.name = u'Times New Roman'
# 设置中文字体
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'方正小标宋_GBK')
run.font.size = Pt(22)
run.bold = False





# 保存并打开word
fname = '哈哈.docx'
doc.save(fname)
cmd = os.path.join('start', 'WINWORD.EXE', os.path.join(os.getcwd(), fname))
subprocess.Popen(cmd, shell=True)
