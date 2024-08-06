# -*- coding: utf-8 -*-

import os,subprocess
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, ns
from docx.shared import Pt, Cm
from docx.enum.text import *
from docx import Document
from docx.text.paragraph import Paragraph
from docx.parts.image import ImagePart



## 总逻辑


# 删除段落
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    # p._p = p._element = None
    paragraph._p = paragraph._element = None


# 判断是否为落款格式
def LuoKuan(str):
    for i in str:
        if i in PUNC:
            return False
    if ((str[0] in NUM) and (str[-1] == "日") and (len(str) <= 12)) or (
        (str[0] in CN_NUM) and (str[-1] == "日") and (len(str) <= 12) or(
         (str[-1] == "部") and (len(str) <= 12 ) )
    ):
        return True
    else:
        return False


def setMargin(docx):
    section = docx.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.9)
    section.right_margin = Cm(2.9)
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(3.0)


# 判断是否为一级标题格式（如：一、xxx）
def GradeOneTitle(str):
    if ((str[0] in CN_NUM) and (str[1] == "、")) or (
        (str[0] in CN_NUM) and (str[1] in CN_NUM) and (str[2] == "、")
    ):
        return True
    else:
        return False


# 判断是否为二级标题格式（如：（一）xxx）
def GradeTwoTitle(str):
    if ((str[0] == "（") and (str[1] in CN_NUM) and (str[2] == "）")) or (
        (str[0] == "（")
        and (str[1] in CN_NUM)
        and (str[2] in CN_NUM)
        and (str[3] == "）")
    ):
        return True
    else:
        return False


# 判断是否为三级标题格式（如：1.xxx）
def GradeThreeTitle(str):
    if ((str[0] in NUM) and (str[1] in PUNC)) or (
        (str[0] in NUM) and (str[1] in NUM) and (str[2] in PUNC)
    ):
        return True
    else:
        return False


# 判断是否为四级标题格式（如：（1）xxx）
def GradeFourTitle(str):
    if ((str[0] == "（") and (str[1] in NUM) and (str[2] == "）")) or (
        (str[0] == "（") and (str[1] in NUM) and (str[2] in NUM) and (str[3] == "）")
    ):
        return True
    else:
        return False


# 判断是否为五级标题格式（如：一是XXX）
def GradeFiveTitle(str):
    if ((str[0] in CN_NUM) and (str[1] in MUST)) or (
        (str[0] in CN_NUM) and (str[1] in CN_NUM) and (str[1] in MUST)
    ):
        return True
    else:
        return False
    
       

# parameters
CN_NUM = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
NUM = ["1", "2", "3", "4", "5", "6", "7", "8", "9", "0"]
PUNC = ["。", "，", "！", "？", "：", "；", "、", ".", "（", "）", "．"]
MUST = ["要", "是", "能"]

def sytle_normalization(filename):
    doc = Document(filename)

    paragraphcnt = 0
    for paragraph in doc.paragraphs:
        paragraphcnt = paragraphcnt + 1

        # 如果该段包含图片，不做任何处理，保持原样
        if any("pic:pic" in run.element.xml for run in paragraph.runs):
            continue
        
        # 段落处理
        paragraph.text = paragraph.text.replace(",", "，")
        paragraph.text = paragraph.text.replace(";", "；")
        paragraph.text = paragraph.text.replace(":", "：")
        paragraph.text = paragraph.text.replace("!", "！")
        paragraph.text = paragraph.text.replace("?", "？")
        paragraph.text = paragraph.text.replace("(", "（")
        paragraph.text = paragraph.text.replace(")", "）")
        paragraph.text = paragraph.text.replace(" ", "")
        paragraph.text = paragraph.text.replace("\t", "")
        paragraph.text = paragraph.text.replace("\n", "")
        if paragraph.text == "":
            delete_paragraph(paragraph)
            paragraphcnt = paragraphcnt - 1
            continue
        paragraph.paragraph_format.left_indent = (
            0  # 预先对缩进赋值, 防止对象为空报错
        )
        paragraph.paragraph_format.element.pPr.ind.set(
            qn("w:firstLineChars"), "0"
        )  # 并去除缩进
        paragraph.paragraph_format.element.pPr.ind.set(qn("w:firstLine"), "0")
        paragraph.paragraph_format.element.pPr.ind.set(qn("w:leftChars"), "0")
        paragraph.paragraph_format.element.pPr.ind.set(qn("w:left"), "0")
        paragraph.paragraph_format.element.pPr.ind.set(qn("w:rightChars"), "0")
        paragraph.paragraph_format.element.pPr.ind.set(qn("w:right"), "0")
        print("这是第%s段" % paragraphcnt)
        print(paragraph.text)


        if paragraphcnt == 1 and len(paragraph.text) < 40:
            # 标题（方正小标宋_GBK、2号、加粗、居中、下端按2号字空一行）
            paragraph.paragraph_format.line_spacing = Pt(29)  # 行距固定值29磅
            paragraph.paragraph_format.space_after = Pt(0)  # 段后间距=0
            for run in paragraph.runs:
                run.font.size = Pt(22)  # 字体大小2号
                run.bold = False  # 加粗
                run.font.name = "Times New Roman"  # 控制是西文时的字体
                run.element.rPr.rFonts.set(
                    qn("w:eastAsia"), "方正小标宋_GBK"
                )  # 控制是中文时的字体
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中
            continue
        elif paragraphcnt == 2 and len(paragraph.text) < 30:
            # 作者单位、姓名
            paragraph.paragraph_format.line_spacing = Pt(29)  # 行距固定值29磅
            paragraph.paragraph_format.space_after = Pt(0)  # 段后间距=0
            for run in paragraph.runs:
                run.font.size = Pt(16)  # 字体大小3号
                run.bold = False  # 加粗
                run.font.name = "Times New Roman"  # 控制是西文时的字体
                run.element.rPr.rFonts.set(
                    qn("w:eastAsia"), "方正楷体_GBK"
                )  # 控制是中文时的字体
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中
            continue
        elif (
            paragraphcnt == 3
            and len(paragraph.text) < 30
            and (paragraph.text[0] == "（")
            and (paragraph.text[1] in NUM)
        ):
            # 日期，如（2023年6月15日）
            paragraph.paragraph_format.line_spacing = Pt(29)  # 行距固定值29磅
            paragraph.paragraph_format.space_after = Pt(0)  # 段后间距=0
            for run in paragraph.runs:
                run.font.size = Pt(16)  # 字体大小3号
                run.bold = False  # 加粗
                run.font.name = "Times New Roman"  # 控制是西文时的字体
                run.element.rPr.rFonts.set(
                    qn("w:eastAsia"), "方正楷体_GBK"
                )  # 控制是中文时的字体
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中
            continue
        # # 处理正文，从第四段开始
        else:
            paragraph.paragraph_format.line_spacing = Pt(29)  # 行距固定值29磅
            paragraph.paragraph_format.space_after = Pt(0)  # 段后间距=0
            paragraph.paragraph_format.first_line_indent = Pt(32)
            for run in paragraph.runs:
                run.font.size = Pt(16)  # 字体大小3号
                run.bold = False  # 字体不加粗
                run.font.name = "Times New Roman"
                run.element.rPr.rFonts.set(qn("w:eastAsia"), "方正仿宋_GBK")
                if GradeOneTitle(
                    run.text
                ):  # 判断是否为一级标题格式（如：一、xxx）
                    run.font.name = "Times New Roman"
                    run.element.rPr.rFonts.set(qn("w:eastAsia"), "方正黑体_GBK")
                elif GradeTwoTitle(
                    run.text
                ):  # 判断是否为二级标题格式（如：（一）xxx）
                    if "。" not in run.text and "：" not in run.text:
                        run.font.name = "Times New Roman"
                        run.element.rPr.rFonts.set(qn("w:eastAsia"), "方正楷体_GBK")
                    else:
                        if "：" in run.text and "。" in run.text:
                            if run.text.index("：") < run.text.index("。"):
                                sentence_to_bold = run.text.split("：")[0] + "："
                                sentence_not_to_bold = run.text.split("：", 1)[1]
                        elif "：" in run.text and "。" not in run.text:
                            sentence_to_bold = run.text.split("：")[0] + "："
                            sentence_not_to_bold = run.text.split("：")[1]
                        elif "。" in run.text:
                            sentence_to_bold = run.text.split("。")[0] + "。"
                            sentence_not_to_bold = run.text.split("。")[1]
                        else:
                            continue                            
                        paragraph.insert_paragraph_before(sentence_to_bold)
                        doc.paragraphs[
                            paragraphcnt - 1
                        ].paragraph_format.first_line_indent = Pt(32)
                        doc.paragraphs[
                            paragraphcnt - 1
                        ].paragraph_format.line_spacing = Pt(
                            29
                        )  # 行距固定值29磅
                        doc.paragraphs[
                            paragraphcnt - 1
                        ].paragraph_format.space_after = Pt(
                            0
                        )  # 段后间距=0
                        doc.paragraphs[paragraphcnt - 1].runs[
                            0
                        ].font.name = "Times New Roman"
                        doc.paragraphs[paragraphcnt - 1].runs[0].font.size = (
                            Pt(16)
                        )  # 字体大小3号
                        doc.paragraphs[paragraphcnt - 1].runs[
                            0
                        ].element.rPr.rFonts.set(
                            qn("w:eastAsia"), "方正仿宋_GBK"
                        )
                        doc.paragraphs[paragraphcnt - 1].runs[
                            0
                        ].bold = True  # 字体加粗
                        doc.paragraphs[paragraphcnt - 1].add_run(
                            sentence_not_to_bold
                        ).bold = False
                        doc.paragraphs[paragraphcnt - 1].runs[
                            1
                        ].font.name = "Times New Roman"
                        doc.paragraphs[paragraphcnt - 1].runs[1].font.size = (
                            Pt(16)
                        )  # 字体大小3号
                        doc.paragraphs[paragraphcnt - 1].runs[
                            1
                        ].element.rPr.rFonts.set(
                            qn("w:eastAsia"), "方正仿宋_GBK"
                        )
                        delete_paragraph(paragraph)
                        
                elif GradeThreeTitle(
                    run.text
                ):  # 判断是否为三级标题格式（如：1.xxx）
                    run.font.name = "Times New Roman"
                    run.element.rPr.rFonts.set(
                            qn("w:eastAsia"), "方正仿宋_GBK"
                        )
                    if "。" not in run.text and "；" not in run.text and len(run.text)<16:
                        run.bold = True  # 字体加粗
                elif GradeFourTitle(
                    run.text
                ):  # 判断是否为四级标题格式（如：（1）xxx）
                    run.font.name = "Times New Roman"
                    run.element.rPr.rFonts.set(
                            qn("w:eastAsia"), "方正仿宋_GBK"
                    )
                #     if "。" not in run.text:
                #         run.font.name = "Times New Roman"
                #         run.element.rPr.rFonts.set(
                #             qn("w:eastAsia"), "方正仿宋_GBK"
                #         )
                #         run.bold = True  # 字体加粗
                #     else:
                #         sentence_to_bold = run.text.split("。")[0] + "。"
                #         sentence_not_to_bold = run.text.split("。", 1)[1]
                #         paragraph.insert_paragraph_before(sentence_to_bold)
                #         doc.paragraphs[
                #             paragraphcnt - 1
                #         ].paragraph_format.first_line_indent = Pt(32)
                #         doc.paragraphs[
                #             paragraphcnt - 1
                #         ].paragraph_format.line_spacing = Pt(
                #             29
                #         )  # 行距固定值29磅
                #         doc.paragraphs[
                #             paragraphcnt - 1
                #         ].paragraph_format.space_after = Pt(
                #             0
                #         )  # 段后间距=0
                #         doc.paragraphs[paragraphcnt - 1].runs[
                #             0
                #         ].font.name = "Times New Roman"
                #         doc.paragraphs[paragraphcnt - 1].runs[0].font.size = (
                #             Pt(16)
                #         )  # 字体大小3号
                #         doc.paragraphs[paragraphcnt - 1].runs[
                #             0
                #         ].element.rPr.rFonts.set(
                #             qn("w:eastAsia"), "方正仿宋_GBK"
                #         )
                #         doc.paragraphs[paragraphcnt - 1].runs[
                #             0
                #         ].bold = True  # 字体加粗
                #         doc.paragraphs[paragraphcnt - 1].add_run(
                #             sentence_not_to_bold
                #         ).bold = False
                #         doc.paragraphs[paragraphcnt - 1].runs[
                #             1
                #         ].font.name = "Times New Roman"
                #         doc.paragraphs[paragraphcnt - 1].runs[1].font.size = (
                #             Pt(16)
                #         )  # 字体大小3号
                #         doc.paragraphs[paragraphcnt - 1].runs[
                #             1
                #         ].element.rPr.rFonts.set(
                #             qn("w:eastAsia"), "方正仿宋_GBK"
                #         )
                #         delete_paragraph(paragraph)
                # elif GradeFiveTitle(
                #     run.text
                # ):  # 判断是否为五级标题格式（如：一是xxx）
                #     if "。" not in run.text:
                #         run.font.name = "Times New Roman"
                #         run.element.rPr.rFonts.set(
                #             qn("w:eastAsia"), "方正仿宋_GBK"
                #         )
                #         run.bold = True  # 字体加粗
                #     else:
                #         sentence_to_bold = run.text.split("。")[0] + "。"
                #         sentence_not_to_bold = run.text.split("。", 1)[1]
                #         paragraph.insert_paragraph_before(sentence_to_bold)
                #         doc.paragraphs[
                #             paragraphcnt - 1
                #         ].paragraph_format.first_line_indent = Pt(32)
                #         doc.paragraphs[
                #             paragraphcnt - 1
                #         ].paragraph_format.line_spacing = Pt(
                #             29
                #         )  # 行距固定值29磅
                #         doc.paragraphs[
                #             paragraphcnt - 1
                #         ].paragraph_format.space_after = Pt(
                #             0
                #         )  # 段后间距=0
                #         doc.paragraphs[paragraphcnt - 1].runs[
                #             0
                #         ].font.name = "Times New Roman"
                #         doc.paragraphs[paragraphcnt - 1].runs[0].font.size = (
                #             Pt(16)
                #         )  # 字体大小3号
                #         doc.paragraphs[paragraphcnt - 1].runs[
                #             0
                #         ].element.rPr.rFonts.set(
                #             qn("w:eastAsia"), "方正仿宋_GBK"
                #         )
                #         doc.paragraphs[paragraphcnt - 1].runs[
                #             0
                #         ].bold = True  # 字体加粗
                #         doc.paragraphs[paragraphcnt - 1].add_run(
                #             sentence_not_to_bold
                #         ).bold = False
                #         doc.paragraphs[paragraphcnt - 1].runs[
                #             1
                #         ].font.name = "Times New Roman"
                #         doc.paragraphs[paragraphcnt - 1].runs[1].font.size = (
                #             Pt(16)
                #         )  # 字体大小3号
                #         doc.paragraphs[paragraphcnt - 1].runs[
                #             1
                #         ].element.rPr.rFonts.set(
                #             qn("w:eastAsia"), "方正仿宋_GBK"
                #         )
                #         delete_paragraph(paragraph)
                elif LuoKuan(run.text):  # 判断是否为落款格式
                    run.font.name = "Times New Roman"
                    run.element.rPr.rFonts.set(qn("w:eastAsia"), "方正仿宋_GBK")
                    run.text = (
                        "\r" * 2 + run.text
                    )  # 前置空格，顶到最右，需手动调整空格
                    paragraph.paragraph_format.left_indent = Pt(
                        288
                    )  # 18B*16Pt=288Pt
                else:  # 普通正文格式
                    run.font.name = "Times New Roman"
                    run.element.rPr.rFonts.set(qn("w:eastAsia"), "方正仿宋_GBK")

    paragraphcnt = 0
    for paragraph in doc.paragraphs:
        paragraphcnt = paragraphcnt + 1
        if paragraphcnt == 1 and len(paragraph.text) < 40:
            run = paragraph.add_run('\n') #空行
            run.font.size = Pt(16)
            run.font.name = "方正楷体_GBK"
            continue
        elif paragraphcnt == 2 and len(paragraph.text) < 30:
            continue
        elif paragraphcnt == 3:
            # run = paragraph.add_run('\n') //空行
            run.font.size = Pt(16)
            run.font.name = "方正楷体_GBK"
            continue

    setMargin(doc)
    return doc



def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(paragraph):

    pre_run = paragraph.add_run()
    t2 = create_element('w:t')
    create_attribute(t2, 'xml:space', 'preserve')
    t2.text = "- "
    pre_run._r.append(t2)
    pre_run.font.size = Pt(12)
    pre_run.font.name = "宋体"

    page_num_run = paragraph.add_run()
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)

    page_num_run.font.size = Pt(12)
    page_num_run.font.name = "宋体"

    post_run = paragraph.add_run()
    t3 = create_element('w:t')
    create_attribute(t3, 'xml:space', 'preserve')
    t3.text = " -"
    post_run._r.append(t3)
    post_run.font.size = Pt(12)
    post_run.font.name = "宋体"
    return


if __name__=="__main__":
    # 测试文档，无格式
    draft_doc = "扬州市数据资产登记指引（试行）_0614V3(1).docx"
    # draft_doc = "draft.docx"
    doc = sytle_normalization(draft_doc)
    # 保存格式化的文档
    fname = 'formatted.docx'
    add_page_number(doc.sections[0].footer.paragraphs[0])
    doc.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.save(fname)

    # 打开格式化后的文档
    cmd = os.path.join('start', 'WINWORD.EXE', os.path.join(os.getcwd(), fname))
    subprocess.Popen(cmd, shell=True)
